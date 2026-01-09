import unittest
import io
import sys
import os
import pandas as pd
import docx

# Add parent dir to path to import logic
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from indexer_logic import extract_word, extract_excel, clean_text

class TestIndexerLogic(unittest.TestCase):

    def test_clean_text(self):
        self.assertEqual(clean_text("  hello   world  "), "hello world")
        self.assertEqual(clean_text(""), "")

    def test_extract_word(self):
        # Create a real docx in memory
        buffer = io.BytesIO()
        doc = docx.Document()
        doc.core_properties.title = "Test Title"
        doc.add_heading("Heading 1", level=1)
        doc.add_paragraph("Hello world")
        doc.save(buffer)
        buffer.seek(0)
        
        content = extract_word(buffer.getvalue())
        self.assertIn("# Test Title", content)
        self.assertIn("## Heading 1", content)
        self.assertIn("Hello world", content)

    def test_extract_excel(self):
        # Create a real excel in memory
        buffer = io.BytesIO()
        df = pd.DataFrame({'Col1': [1, 2], 'Col2': ['A', 'B']})
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Sheet1')
        buffer.seek(0)
        
        content = extract_excel(buffer.getvalue())
        self.assertIn("## Sheet: Sheet1", content)
        # Check for column headers and data existence rather than exact spacing
        self.assertIn("Col1", content)
        self.assertIn("Col2", content) 
        self.assertIn("A ", content.replace('|', ' ')) # Simple check

if __name__ == '__main__':
    unittest.main()
